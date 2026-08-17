#!/usr/bin/env python3
"""Build the polished Arabic owner-review DOCX from the simplified V5 Markdown."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DEFAULT_INPUT = ROOT / "TAQNYAT_INVITATION_TEMPLATES_OWNER_REVIEW_V5.md"
DEFAULT_OUTPUT = ROOT / "TAQNYAT_INVITATION_TEMPLATES_OWNER_SHORTLIST_V5.docx"

# Source template numbers selected for the owner-review shortlist.
# Wedding keeps four distinct voices; every other category keeps two.
SHORTLIST = {
    1: (1, 2, 3, 5),  # الزفاف
    2: (1, 3),        # الخطوبة
    3: (2, 4),        # عيد الميلاد
    4: (1, 2),        # استقبال المولود
    5: (1, 4),        # المناسبة النسائية
    6: (1, 5),        # المناسبة العامة
    7: (1, 2),        # المؤتمر والفعاليات المهنية
}

FONT = "Tahoma"
NAVY = "17324D"
GOLD = "B68A4A"
INK = "26323C"
MUTED = "66727C"
SAND = "F4EEE4"
IVORY = "FBF9F5"
LINE = "DED6C8"
WHITE = "FFFFFF"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_font(run, size: float, *, bold: bool = False, color: str = INK, font: str = FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = False
    run.font.color.rgb = rgb(color)
    rpr = run._r.get_or_add_rPr()
    italic_cs = rpr.find(qn("w:iCs"))
    if italic_cs is None:
        italic_cs = OxmlElement("w:iCs")
        rpr.append(italic_cs)
    italic_cs.set(qn("w:val"), "0")
    rfonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), font)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "ar-SA")
    lang.set(qn("w:bidi"), "ar-SA")
    return run


def set_run_rtl(run, enabled: bool = True):
    rpr = run._r.get_or_add_rPr()
    rtl = rpr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rpr.append(rtl)
    rtl.set(qn("w:val"), "1" if enabled else "0")


def set_rtl(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    # Word mirrors left/right alignment after bidi is applied. Supplying the
    # opposite physical value keeps Arabic paragraphs visually right-aligned.
    if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif alignment == WD_ALIGN_PARAGRAPH.LEFT:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = alignment
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1")
    return paragraph


def shade_paragraph(paragraph, fill: str):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def shade_run(run, fill: str):
    rpr = run._r.get_or_add_rPr()
    shd = rpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        rpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_border(paragraph, *, side: str, color: str, size: int = 12, space: int = 8):
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    edge = borders.find(qn(f"w:{side}"))
    if edge is None:
        edge = OxmlElement(f"w:{side}")
        borders.append(edge)
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)


def set_spacing(paragraph, *, before=0, after=0, line=1.2):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(paragraph, text: str, size: float, *, bold=False, color=INK, rtl=True):
    run = paragraph.add_run(text)
    set_font(run, size, bold=bold, color=color)
    set_run_rtl(run, rtl)
    return run


def arabic_number(value: int) -> str:
    return str(value).translate(str.maketrans("0123456789", "٠١٢٣٤٥٦٧٨٩"))


def parse_catalog(path: Path):
    lines = path.read_text(encoding="utf-8").splitlines()
    category_re = re.compile(r"^##\s+(\d+)\.\s+(.+)$")
    template_re = re.compile(r"^###\s+(\d+)\.(\d+)\s+(.+)$")
    categories = []
    current_category = None
    current_template = None
    in_code = False
    body_lines = []

    for line in lines:
        category_match = category_re.match(line)
        if category_match and not in_code:
            current_category = {
                "number": int(category_match.group(1)),
                "name": category_match.group(2).strip(),
                "templates": [],
            }
            categories.append(current_category)
            current_template = None
            continue

        template_match = template_re.match(line)
        if template_match and not in_code:
            if current_category is None:
                raise ValueError("A template appeared before its category.")
            current_template = {
                "number": int(template_match.group(2)),
                "title": template_match.group(3).strip(),
                "body": "",
            }
            current_category["templates"].append(current_template)
            continue

        if line == "```text" and current_template is not None:
            in_code = True
            body_lines = []
            continue
        if line == "```" and in_code:
            current_template["body"] = "\n".join(body_lines).strip()
            in_code = False
            continue
        if in_code:
            body_lines.append(line)

    template_count = sum(len(category["templates"]) for category in categories)
    if len(categories) != 7 or template_count != 35:
        raise ValueError(f"Expected 7 categories and 35 templates; found {len(categories)} and {template_count}.")
    if any(len(category["templates"]) != 5 for category in categories):
        raise ValueError("Every category must contain exactly five templates.")
    if any(not template["body"] for category in categories for template in category["templates"]):
        raise ValueError("Every template must have a text body.")
    return categories


def apply_shortlist(categories):
    shortlisted = []
    for category in categories:
        selected_numbers = SHORTLIST[category["number"]]
        selected = [
            template.copy()
            for template in category["templates"]
            if template["number"] in selected_numbers
        ]
        if [template["number"] for template in selected] != list(selected_numbers):
            raise ValueError(f"Shortlist mismatch in category {category['number']}.")
        for display_number, template in enumerate(selected, start=1):
            template["source_number"] = template["number"]
            template["number"] = display_number
        shortlisted.append({**category, "templates": selected})

    expected_counts = [4, 2, 2, 2, 2, 2, 2]
    actual_counts = [len(category["templates"]) for category in shortlisted]
    if actual_counts != expected_counts:
        raise ValueError(f"Expected shortlist counts {expected_counts}; found {actual_counts}.")
    return shortlisted


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(17)
    section.left_margin = Mm(19)
    section.right_margin = Mm(19)
    section.header_distance = Mm(9)
    section.footer_distance = Mm(9)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, bold in (
        ("Title", 30, NAVY, True),
        ("Subtitle", 13, MUTED, False),
        ("Heading 1", 23, NAVY, True),
        ("Heading 2", 13.5, NAVY, True),
    ):
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = False
        style.font.color.rgb = rgb(color)
        rpr = style.element.get_or_add_rPr()
        italic_cs = rpr.find(qn("w:iCs"))
        if italic_cs is None:
            italic_cs = OxmlElement("w:iCs")
            rpr.append(italic_cs)
        italic_cs.set(qn("w:val"), "0")
        rfonts = rpr.get_or_add_rFonts()
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{key}"), FONT)
        ppr = style.element.get_or_add_pPr()
        inherited_borders = ppr.find(qn("w:pBdr"))
        if inherited_borders is not None:
            ppr.remove(inherited_borders)

    if "Template Body" not in styles:
        template_style = styles.add_style("Template Body", WD_STYLE_TYPE.PARAGRAPH)
    else:
        template_style = styles["Template Body"]
    template_style.font.name = FONT
    template_style.font.size = Pt(10.5)
    template_style.font.color.rgb = rgb(INK)
    template_style.paragraph_format.left_indent = Mm(5)
    template_style.paragraph_format.right_indent = Mm(5)
    template_style.paragraph_format.line_spacing = 1.35

    doc.core_properties.title = "مختارات مراجعة مالك المشروع — نصوص دعوات هلا V5"
    doc.core_properties.subject = "16 نصًا مختارًا ضمن سبع فئات"
    doc.core_properties.author = "هلا"
    doc.core_properties.keywords = "هلا، دعوات، تقنيات، ميتا، مراجعة نصوص"


def add_page_field(paragraph):
    set_rtl(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, result, end):
        run._r.append(node)
    set_font(run, 8.5, color=MUTED)
    set_run_rtl(run, False)


def add_header_footer(doc: Document):
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    set_rtl(header_p, WD_ALIGN_PARAGRAPH.RIGHT)
    set_spacing(header_p, after=2)
    add_text(header_p, "هلا  |  مراجعة نصوص الدعوات", 8.5, bold=True, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    add_page_field(footer_p)

    first_header = section.first_page_header.paragraphs[0]
    first_header.clear()
    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.clear()


def add_cover(doc: Document, categories):
    spacer = doc.add_paragraph()
    set_spacing(spacer, after=34)

    brand = doc.add_paragraph()
    set_rtl(brand)
    set_spacing(brand, after=18)
    add_text(brand, "هــــلا", 13, bold=True, color=GOLD)

    title = doc.add_paragraph(style="Title")
    set_rtl(title)
    set_spacing(title, after=8, line=1.05)
    add_text(title, "مختارات نصوص الدعوات", 30, bold=True, color=NAVY)

    subtitle = doc.add_paragraph(style="Subtitle")
    set_rtl(subtitle)
    set_spacing(subtitle, after=24)
    add_text(subtitle, "نسخة مختصرة لمراجعة مالك المشروع", 14, color=MUTED)

    summary = doc.add_paragraph()
    set_rtl(summary)
    summary.paragraph_format.left_indent = Mm(6)
    summary.paragraph_format.right_indent = Mm(6)
    set_spacing(summary, before=4, after=22, line=1.45)
    shade_paragraph(summary, SAND)
    set_border(summary, side="right", color=GOLD, size=18, space=10)
    add_text(
        summary,
        "أربعة نصوص مختارة للزفاف، ونصّان لكل فئة أخرى، بنوع تشغيل واحد لتسهيل المراجعة قبل الاعتماد.",
        11.5,
        color=INK,
    )

    stats = doc.add_paragraph()
    set_rtl(stats, WD_ALIGN_PARAGRAPH.CENTER)
    set_spacing(stats, before=10, after=8)
    add_text(stats, "١٦ نصًا   •   ٧ فئات   •   نوع واحد للمراجعة", 11, bold=True, color=NAVY)

    mode = doc.add_paragraph()
    set_rtl(mode, WD_ALIGN_PARAGRAPH.CENTER)
    set_spacing(mode, after=36)
    add_text(mode, "ردّ على الدعوة مع إرسال رمز الدخول بعد تأكيد الحضور", 9.5, color=MUTED)

    edition = doc.add_paragraph()
    set_rtl(edition, WD_ALIGN_PARAGRAPH.CENTER)
    set_spacing(edition, before=18)
    add_text(edition, "الإصدار الخامس  |  للمراجعة الداخلية", 8.5, color=GOLD)
    doc.add_page_break()


def add_contents(doc: Document, categories):
    eyebrow = doc.add_paragraph()
    set_rtl(eyebrow)
    set_spacing(eyebrow, after=5)
    add_text(eyebrow, "دليل الملف", 9, bold=True, color=GOLD)

    heading = doc.add_paragraph(style="Heading 1")
    set_rtl(heading)
    set_spacing(heading, after=8)
    add_text(heading, "الفئات المشمولة", 23, bold=True, color=NAVY)

    intro = doc.add_paragraph()
    set_rtl(intro)
    set_spacing(intro, after=20, line=1.4)
    add_text(intro, "رتّبنا المختارات حسب المناسبة. تبدأ كل فئة في صفحة مستقلة، مع أربعة خيارات للزفاف وخيارين لكل فئة أخرى.", 10.5, color=MUTED)

    for category in categories:
        row = doc.add_paragraph()
        set_rtl(row)
        row.paragraph_format.left_indent = Mm(3)
        row.paragraph_format.right_indent = Mm(3)
        set_spacing(row, before=3, after=3, line=1.3)
        set_border(row, side="bottom", color=LINE, size=5, space=7)
        add_text(row, f"{arabic_number(category['number'])}  {category['name']}", 12, bold=True, color=NAVY)
        count_label = "أربعة نصوص" if len(category["templates"]) == 4 else "نصّان"
        add_text(row, f"    {count_label}", 9, color=MUTED)

    doc.add_page_break()


PLACEHOLDER_RE = re.compile(r"(\{\{\d+\}\})")


def add_rich_line(paragraph, text: str):
    for part in PLACEHOLDER_RE.split(text):
        if not part:
            continue
        if PLACEHOLDER_RE.fullmatch(part):
            run = add_text(paragraph, part, 9.5, bold=True, color=NAVY, rtl=False)
            shade_run(run, "EDE2CF")
        else:
            add_text(paragraph, part, 10.5, color=INK)


def add_template(doc: Document, template):
    title = doc.add_paragraph(style="Heading 2")
    set_rtl(title)
    title.paragraph_format.keep_with_next = True
    set_spacing(title, before=11, after=6, line=1.2)
    set_border(title, side="right", color=GOLD, size=14, space=9)
    add_text(title, f"{arabic_number(template['number'])}  {template['title']}", 13.5, bold=True, color=NAVY)

    groups = [group.splitlines() for group in re.split(r"\n\s*\n", template["body"]) if group.strip()]
    for group_index, lines in enumerate(groups):
        paragraph = doc.add_paragraph(style="Template Body")
        set_rtl(paragraph)
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = group_index < len(groups) - 1
        set_spacing(
            paragraph,
            before=5 if group_index == 0 else 0,
            after=6 if group_index == len(groups) - 1 else 0,
            line=1.35,
        )
        shade_paragraph(paragraph, IVORY)
        set_border(paragraph, side="right", color=LINE, size=6, space=8)
        for line_index, line in enumerate(lines):
            add_rich_line(paragraph, line)
            if line_index < len(lines) - 1:
                paragraph.add_run().add_break()


def add_category(doc: Document, category):
    category_no = doc.add_paragraph()
    set_rtl(category_no)
    set_spacing(category_no, after=2)
    add_text(category_no, f"الفئة {arabic_number(category['number'])}", 9, bold=True, color=GOLD)

    heading = doc.add_paragraph(style="Heading 1")
    set_rtl(heading)
    set_spacing(heading, after=5, line=1.1)
    add_text(heading, category["name"], 23, bold=True, color=NAVY)

    subheading = doc.add_paragraph()
    set_rtl(subheading)
    set_spacing(subheading, after=13)
    set_border(subheading, side="bottom", color=GOLD, size=8, space=8)
    count_label = "أربعة نصوص مختارة" if len(category["templates"]) == 4 else "نصّان مختاران"
    add_text(subheading, count_label, 9.5, color=MUTED)

    for template in category["templates"]:
        add_template(doc, template)


def build(input_path: Path, output_path: Path):
    categories = apply_shortlist(parse_catalog(input_path))
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_cover(doc, categories)
    add_contents(doc, categories)

    for index, category in enumerate(categories):
        if index:
            doc.add_page_break()
        add_category(doc, category)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return len(categories), sum(len(category["templates"]) for category in categories)


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT, help="Simplified V5 Markdown source")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT, help="Destination .docx file")
    args = parser.parse_args()
    categories, templates = build(args.input.resolve(), args.output.resolve())
    print(f"Created {args.output.resolve()} ({categories} categories, {templates} templates).")


if __name__ == "__main__":
    main()
